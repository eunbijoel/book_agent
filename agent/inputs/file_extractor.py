from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

from agent.inputs.base import BaseExtractor

logger = logging.getLogger(__name__)

_MAX_FILE_SIZE = 100 * 1024 * 1024  # 100 MB
_ENCODINGS = ("utf-8", "cp949", "latin-1")
_SUPPORTED = {".txt", ".md", ".pdf", ".docx", ".html", ".htm"}
_REMOVE_TAGS = {"script", "style", "nav", "header", "footer", "aside", "noscript", "iframe"}


class FileExtractor(BaseExtractor):
    """Extract text from local files (TXT, MD, PDF, DOCX)."""

    def extract(self, source: str) -> dict[str, Any]:
        path = Path(source)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {source}")

        if path.stat().st_size > _MAX_FILE_SIZE:
            raise ValueError(f"File too large (>100 MB): {source}")

        suffix = path.suffix.lower()
        if suffix not in _SUPPORTED:
            raise ValueError(f"Unsupported format: {suffix}. Supported: {', '.join(sorted(_SUPPORTED))}")

        if suffix in (".txt", ".md"):
            text, title = self._extract_text(path)
        elif suffix == ".pdf":
            text, title = self._extract_pdf(path)
        elif suffix == ".docx":
            text, title = self._extract_docx(path)
        elif suffix in (".html", ".htm"):
            text, title = self._extract_html(path)
        else:
            raise ValueError(f"Unsupported format: {suffix}")

        if len(text) < 100:
            logger.warning("Extracted text very short (%d chars) from %s", len(text), source)

        chunks = self._chunk_text(text)

        return {
            "source_type": suffix.lstrip("."),
            "source_path": str(path),
            "title": title or path.stem,
            "content": text[:50000],
            "content_length": len(text),
            "chunks": chunks,
            "metadata": {
                "file_name": path.name,
                "file_size": path.stat().st_size,
            },
        }

    def _extract_text(self, path: Path) -> tuple[str, str]:
        text = self._read_with_fallback(path)
        title = ""
        for line in text.splitlines():
            stripped = line.strip()
            if stripped:
                if stripped.startswith("# "):
                    title = stripped[2:].strip()
                else:
                    title = stripped[:100]
                break
        return text, title

    def _extract_pdf(self, path: Path) -> tuple[str, str]:
        import pdfplumber

        pages: list[str] = []
        title = ""
        with pdfplumber.open(path) as pdf:
            if pdf.metadata and pdf.metadata.get("Title"):
                title = pdf.metadata["Title"]

            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pages.append(page_text)

                for table in page.extract_tables():
                    rows = []
                    for row in table:
                        cells = [str(c) if c is not None else "" for c in row]
                        rows.append(" | ".join(cells))
                    if rows:
                        pages.append("\n".join(rows))

        text = "\n\n".join(pages)
        return text, title

    def _extract_docx(self, path: Path) -> tuple[str, str]:
        from docx import Document

        doc = Document(str(path))

        title = ""
        if doc.core_properties.title:
            title = doc.core_properties.title

        paragraphs: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                paragraphs.append("\n".join(rows))

        text = "\n\n".join(paragraphs)
        if not title and paragraphs:
            title = paragraphs[0][:100]
        return text, title

    def _extract_html(self, path: Path) -> tuple[str, str]:
        from bs4 import BeautifulSoup

        raw = self._read_with_fallback(path)
        soup = BeautifulSoup(raw, "html.parser")

        title = ""
        og = soup.find("meta", property="og:title")
        if og and og.get("content"):
            title = og["content"].strip()
        elif soup.find("title") and soup.find("title").string:
            title = soup.find("title").string.strip()
        else:
            h1 = soup.find("h1")
            if h1:
                title = h1.get_text(strip=True)

        for tag in soup.find_all(list(_REMOVE_TAGS)):
            tag.decompose()

        body = None
        for tag_name in ("article", "main"):
            body = soup.find(tag_name)
            if body:
                break
        if not body:
            body = soup.find("body") or soup

        text = body.get_text(separator="\n")
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[ \t]{2,}", " ", text)
        text = text.strip()

        return text, title

    def _read_with_fallback(self, path: Path) -> str:
        for enc in _ENCODINGS:
            try:
                return path.read_text(encoding=enc)
            except (UnicodeDecodeError, UnicodeError):
                continue
        return path.read_text(encoding="utf-8", errors="replace")
